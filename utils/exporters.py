from io import BytesIO
from pathlib import Path
from typing import List
import re
import unicodedata

def _find_logo_file() -> Path | None:
    repo_root = Path(__file__).resolve().parents[1]
    search_dirs = [Path.cwd(), repo_root]
    names = ["neogen_logo", "logo"]
    exts  = [".png", ".jpg", ".jpeg", ".webp", ".gif"]
    for d in search_dirs:
        for n in names:
            for ext in exts:
                p = d / "assets" / f"{n}{ext}"
                if p.exists():
                    return p
    return None

def _find_ttf_font() -> Path | None:
    candidates = [
        Path("assets/fonts/DejaVuSans.ttf"),
        Path("assets/DejaVuSans.ttf"),
        Path(__file__).resolve().parents[1] / "assets" / "fonts" / "DejaVuSans.ttf",
    ]
    for c in candidates:
        if c.exists():
            return c
    return None

_md_h1 = re.compile(r"^#\s+(.*)")
_md_h2 = re.compile(r"^##\s+(.*)")
_md_bullet = re.compile(r"^[-*•]\s+(.*)")

def _parse_markdown(md: str):
    lines = [ln.rstrip() for ln in md.splitlines()]
    buf: List[str] = []
    bullets: List[str] = []
    def flush_par():
        nonlocal buf
        if buf:
            yield ("p", " ".join(buf).strip())
            buf = []
    def flush_ul():
        nonlocal bullets
        if bullets:
            yield ("ul", bullets[:])
            bullets = []
    i = 0
    while i < len(lines):
        ln = lines[i]
        if not ln.strip():
            for item in flush_par(): yield item
            for item in flush_ul(): yield item
            i += 1; continue
        m1 = _md_h1.match(ln)
        m2 = _md_h2.match(ln)
        mb = _md_bullet.match(ln)
        if m1:
            for item in flush_par(): yield item
            for item in flush_ul(): yield item
            yield ("h1", m1.group(1).strip())
        elif m2:
            for item in flush_par(): yield item
            for item in flush_ul(): yield item
            yield ("h2", m2.group(1).strip())
        elif mb:
            for item in flush_par(): yield item
            bullets.append(mb.group(1).strip())
        else:
            buf.append(ln.strip())
        i += 1
    for item in flush_par(): yield item
    for item in flush_ul(): yield item

# ---------------- DOCX (unchanged) ----------------
def markdown_to_docx_bytes(md: str, filename_title: str = "Job Description") -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.core_properties.title = filename_title

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    logo = _find_logo_file()
    if logo and logo.exists():
        hdr = doc.sections[0].header
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(str(logo), height=Inches(0.45))
        except Exception:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for kind, content in _parse_markdown(md):
        if kind == "h1":
            doc.add_heading(content, level=1)
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "p":
            para = doc.add_paragraph(content)
            para.paragraph_format.space_after = Pt(6)
        elif kind == "ul":
            for item in content:
                doc.add_paragraph(item, style="List Bullet")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------------- PDF (fpdf2) with Unicode font or safe fallback ----------------
def _ascii_safe(text: str) -> str:
    # Replace common unicode punctuation with ASCII so core fonts can render
    repl = {
        "•": "-", "–": "-", "—": "-", "-": "-",  # bullets/dashes
        "“": '"', "”": '"', "„": '"', "«": '"', "»": '"',
        "’": "'", "‘": "'", "´": "'", "`": "'",
        "…": "...", "\u00A0": " ", "\u200B": "", "\u2011": "-",  # nbsp/zero-width/no-break
    }
    for k, v in repl.items():
        text = text.replace(k, v)
    # Strip diacritics then encode to latin-1
    text = unicodedata.normalize("NFKD", text)
    return text.encode("latin-1", "replace").decode("latin-1")

def markdown_to_pdf_bytes(md: str, filename_title: str = "Job Description") -> bytes:
    from fpdf import FPDF

    pdf = FPDF(unit="pt", format="A4")
    pdf.set_auto_page_break(auto=True, margin=36)
    pdf.add_page()
    left, top, right = 36, 36, 36
    pdf.set_margins(left, top, right)

    # Try Unicode TTF
    unicode_font = False
    ttf = _find_ttf_font()
    if ttf:
        try:
            pdf.add_font("DejaVu", "", str(ttf), uni=True)
            pdf.set_font("DejaVu", "", 11)
            unicode_font = True
        except Exception:
            pass
    if not unicode_font:
        pdf.set_font("Helvetica", "", 11)

    # Header logo (right)
    logo = _find_logo_file()
    if logo and logo.exists():
        try:
            pdf.image(str(logo), x=pdf.w - right - 120, y=top-6, h=36)
        except Exception:
            pass

    def write_para(text: str, size=11, style=""):
        if unicode_font:
            pdf.set_font("DejaVu", style, size)
            out = text
        else:
            pdf.set_font("Helvetica", style, size)
            out = _ascii_safe(text)
        pdf.multi_cell(w=pdf.w - left - right, h=16, txt=out)
        pdf.ln(2)

    for kind, content in _parse_markdown(md):
        if kind == "h1":
            write_para(content, size=18, style="B")
        elif kind == "h2":
            write_para(content, size=14, style="B")
        elif kind == "p":
            write_para(content, size=11)
        elif kind == "ul":
            bullet = "• " if unicode_font else "- "
            for item in content:
                write_para(bullet + item, size=11)

    out = BytesIO()
    out.write(bytes(pdf.output(dest="S")))
    return out.getvalue()
# ========= Neogen Interview Guide DOCX (boxed layout) =========
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

def _add_logo_header_footer(doc: Document, logo_path: str, footer_text: str = "Powered by Neogen HR"):
    sec = doc.sections[0]
    # margins a bit airier
    sec.top_margin, sec.bottom_margin = Inches(0.6), Inches(0.6)
    sec.left_margin, sec.right_margin = Inches(0.7), Inches(0.7)

    # Header with right-aligned logo
    header = sec.header
    par = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        if Path(logo_path).exists():
            run = par.add_run()
            run.add_picture(logo_path, width=Inches(1.35))
    except Exception:
        pass

    # Footer text, right aligned
    footer = sec.footer
    fpar = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fpar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fpar.text = footer_text
    for r in fpar.runs:
        r.font.size = Pt(8)

def _h_rule(par):
    # poor-man HR: bottom border
    p = par._element
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '9CA3AF')  # slate-400
    pbdr.append(bottom)
    pPr.append(pbdr)

def _section_chip(doc: Document, text: str):
    # heading + thin rule to mimic a "chip"
    h = doc.add_paragraph(text)
    h.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else None
    for r in h.runs: r.font.size = Pt(14)
    _h_rule(doc.add_paragraph(""))

def _card(doc: Document, lines: list[str]):
    # Single-cell table used as a "box"
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    # light border & padding via paragraph spacing
    p = cell.paragraphs[0]
    for ln in lines:
        run = p.add_run(ln + "\n")
        run.font.size = Pt(11)
    # Add a blank paragraph after the box
    doc.add_paragraph("")

def _split_questions(md: str):
    # Very forgiving splitter: looks for lines that start with a question-ish header or bold "Q"
    blocks = []
    cur = []
    for line in md.splitlines():
        if re.match(r'^\s*(#{3,5}\s+|[*-]\s*Q[: ]|Q[: ])', line):
            if cur:
                blocks.append("\n".join(cur).strip())
                cur = []
        cur.append(line)
    if cur: blocks.append("\n".join(cur).strip())
    return blocks

def _parse_block_to_card_lines(block: str):
    # Extract Question / Intent / Good / Follow-ups in a robust, label-insensitive way
    # Accepts headings ####, bold labels, or plain labels with colon.
    text = block.strip()

    # Question line: first non-empty line without the label keywords
    lines = [l.strip(" #*-") for l in text.splitlines() if l.strip()]
    q = next((l for l in lines if re.match(r'(?i)^(Q[:\-\s]|Describe|Tell me|Give an example|Explain|How do you)\b', l)), lines[0] if lines else "")

    def grab(label):
        m = re.search(rf'(?is)(?:^|\n)\s*{label}\s*:\s*(.+?)(?:\n[A-Z][^\n]{0,50}\s*:|\Z)', text, re.IGNORECASE)
        return (m.group(1).strip() if m else "")

    intent = grab(r'(Intent|Purpose|Why)')
    good   = grab(r'(What\s+good\s+looks\s+like|Good\s+looks\s+like|Indicators|Evidence)')
    foll   = grab(r'(Follow[-\s]*ups|Probes|Follow\s*up)')

    out = []
    if q:
        out.append(q if q.endswith("?") else q + "")
        out.append("")
    if intent:
        out.append(f"Intent: {intent}")
    if good:
        out.append(f"What good looks like: {good}")
    if foll:
        out.append(f"Follow-ups: {foll}")
    return out if out else lines

def interview_to_docx_bytes(
    md: str,
    *,
    job_title: str,
    stage: str,
    duration: str = "60 mins",
    logo_path: str = "assets/neogen_logo.png",
    filename_title: str | None = None
) -> bytes:
    """
    Build a Neogen-styled DOCX for interview guides with boxed question cards,
    header logo and footer. We lightly parse common labels from Markdown.
    """
    doc = Document()
    _add_logo_header_footer(doc, logo_path)

    # Cover
    title = doc.add_paragraph(job_title)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in title.runs:
        r.font.size = Pt(22); r.bold = True

    meta = doc.add_paragraph(f"Interview stage: {stage} · Duration: {duration}")
    meta_format = meta.runs[0].font if meta.runs else meta.add_run().font
    meta_format.size = Pt(11)

    doc.add_paragraph("")  # spacing

    # Try to recognize sections by H2 headings
    sections = re.split(r'(?m)^\s*##\s+', md)
    first = sections[0]
    if len(sections) > 1:
        # sections[0] may contain intro; keep scanning remaining
        named = []
        for s in sections[1:]:
            name, _, body = s.partition("\n")
            named.append((name.strip(), body.strip()))
    else:
        named = [("Guide", md)]

    # Render sections
    for name, body in named:
        _section_chip(doc, name)
        # If section appears to be a question set, build cards:
        if re.search(r'(?i)question|core|technical|competenc|closing|culture', name):
            for blk in _split_questions(body):
                lines = _parse_block_to_card_lines(blk)
                if lines: _card(doc, lines)
        else:
            # Otherwise just paragraph content (bullets will carry over into docx as plain lines)
            for para in body.split("\n\n"):
                p = doc.add_paragraph(para.strip())
                for r in p.runs: r.font.size = Pt(11)
            doc.add_paragraph("")

    # Output
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
# ========= Neogen Interview Pack DOCX (template-driven layout) =========
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

# Reuse helpers if present; else define light ones
def _add_logo_header_footer(doc: Document, logo_path: str, footer_text: str = "Powered by Neogen HR"):
    sec = doc.sections[0]
    sec.top_margin, sec.bottom_margin = Inches(0.6), Inches(0.6)
    sec.left_margin, sec.right_margin = Inches(0.7), Inches(0.7)

    header = sec.header
    par = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try:
        if Path(logo_path).exists():
            run = par.add_run()
            run.add_picture(logo_path, width=Inches(1.35))
    except Exception:
        pass

    footer = sec.footer
    fpar = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fpar.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fpar.text = footer_text
    for r in fpar.runs: r.font.size = Pt(8)

def _h_rule(par):
    p = par._element
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '9CA3AF')
    pbdr.append(bottom); pPr.append(pbdr)

def _section_chip(doc: Document, text: str):
    h = doc.add_paragraph(text)
    h.style = doc.styles['Heading 2'] if 'Heading 2' in doc.styles else None
    for r in h.runs: r.font.size = Pt(14); r.bold = True
    _h_rule(doc.add_paragraph(""))

def _bullet_block(doc: Document, text: str):
    # render bullet-ish lines (split by blank line)
    for line in [l.strip() for l in text.splitlines() if l.strip()]:
        p = doc.add_paragraph(line)
        p_format = p.paragraph_format
        p_format.space_after = Pt(2)
        for r in p.runs: r.font.size = Pt(11)
    doc.add_paragraph("")

def _split_questions(md: str):
    # Split on obvious new question markers
    blocks, cur = [], []
    for line in md.splitlines():
        if re.match(r'^\s*(#{3,5}\s+|[*-]\s*Q[: ]|Q[: ]|Tell me|Describe|Give an example|Explain|How do you)\b', line, re.IGNORECASE):
            if cur: blocks.append("\n".join(cur).strip()); cur = []
        cur.append(line)
    if cur: blocks.append("\n".join(cur).strip())
    return blocks

def _parse_block(block: str):
    text = block.strip()
    # First non-empty line is the question unless it starts with a label
    lines = [l.strip(" #*-") for l in text.splitlines() if l.strip()]
    question = lines[0] if lines else ""

    def grab(label):
        m = re.search(rf'(?is)(?:^|\n)\s*{label}\s*:\s*(.+?)(?:\n[A-Z][^\n]{{0,50}}\s*:|\Z)', text, re.IGNORECASE)
        return (m.group(1).strip() if m else "")

    intent = grab(r'(Intent|Purpose|Why)')
    good   = grab(r'(What\s+good\s+looks\s+like|Good\s+looks\s+like|Indicators|Evidence)')
    follow = grab(r'(Follow[-\s]*ups|Probes|Follow\s*up)')
    return question, intent, good, follow

def _card(doc: Document, question: str, intent: str, good: str, follow: str):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    p = cell.paragraphs[0]
    # Question line (bold-ish)
    rq = p.add_run((question or "").strip())
    rq.font.size = Pt(12); rq.bold = True
    p.add_run("\n")

    if intent:
        ri = p.add_run("Intent: "); ri.bold = True; ri.font.size = Pt(11)
        p.add_run(intent + "\n").font.size = Pt(11)
    if good:
        rg = p.add_run("What good looks like: "); rg.bold = True; rg.font.size = Pt(11)
        p.add_run(good + "\n").font.size = Pt(11)
    if follow:
        rf = p.add_run("Follow-ups: "); rf.bold = True; rf.font.size = Pt(11)
        p.add_run(follow).font.size = Pt(11)

    doc.add_paragraph("")

SECTION_ORDER = [
    "Housekeeping",
    "Core Questions",
    "Competency Questions",
    "Technical Questions",
    "Culture & Values",
    "Closing Questions",
    "Close-down & Next Steps",
    "Scoring Rubric",
]

def _extract_section(md: str, name: str) -> str:
    # Grab content under "## {name}" up to the next "##"
    pat = rf'(?is)^\s*##\s*{re.escape(name)}\s*\n(.*?)(?=^\s*##\s|\Z)'
    m = re.search(pat, md, re.MULTILINE)
    return m.group(1).strip() if m else ""

def interview_pack_to_docx_bytes(
    md: str,
    *,
    job_title: str,
    interview_type: str,
    duration: str = "60 mins",
    logo_path: str = "assets/neogen_logo.png",
) -> bytes:
    doc = Document()
    _add_logo_header_footer(doc, logo_path)

    # Cover
    title = doc.add_paragraph(f"{job_title} — Competency Pack")
    for r in title.runs: r.font.size = Pt(22); r.bold = True
    meta = doc.add_paragraph(f"Interview type: {interview_type} · Duration: {duration}")
    if meta.runs: meta.runs[0].font.size = Pt(11)
    doc.add_paragraph("")

    # Iterate sections in fixed order
    for name in SECTION_ORDER:
        body = _extract_section(md, name)
        if not body: 
            continue
        _section_chip(doc, name)

        if name in {"Core Questions","Competency Questions","Technical Questions","Culture & Values","Closing Questions"}:
            # Expect question blocks
            for blk in _split_questions(body):
                q, intent, good, follow = _parse_block(blk)
                if q or intent or good or follow:
                    _card(doc, q, intent, good, follow)
            doc.add_paragraph("")
        elif name in {"Housekeeping","Close-down & Next Steps","Scoring Rubric"}:
            _bullet_block(doc, body)
        else:
            # fallback
            _bullet_block(doc, body)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
